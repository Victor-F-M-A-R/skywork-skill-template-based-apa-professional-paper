from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT = Path(__file__).with_name("apa-professional-paper-template.docx")

FULL_TITLE = "[Full Paper Title in Title Case]"
RUNNING_HEAD = "[SHORTENED TITLE]"
AUTHOR_NAMES = "[Author Name], [Author Name], and [Author Name]"
AFFILIATIONS = "[Department, Institution]; [Department, Institution]"


def set_font(run, italic=False, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic


def set_style_font(style, size=12, bold=False, italic=False):
    font = style.font
    font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)
    set_font(run)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue

        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)

        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.different_first_page_header_footer = False
    section.start_type = WD_SECTION.NEW_PAGE

    normal = document.styles["Normal"]
    set_style_font(normal)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    styles = document.styles

    title = styles.add_style("APA Title", 1)
    title.base_style = normal
    set_style_font(title, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 2
    title.paragraph_format.space_after = Pt(0)

    centered = styles.add_style("APA Centered Bold", 1)
    centered.base_style = normal
    set_style_font(centered, bold=True)
    centered.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    centered.paragraph_format.line_spacing = 2
    centered.paragraph_format.space_after = Pt(0)

    body = styles.add_style("APA Body", 1)
    body.base_style = normal
    set_style_font(body)
    body.paragraph_format.first_line_indent = Inches(0.5)
    body.paragraph_format.line_spacing = 2
    body.paragraph_format.space_after = Pt(0)

    abstract = styles.add_style("APA Abstract", 1)
    abstract.base_style = normal
    set_style_font(abstract)
    abstract.paragraph_format.first_line_indent = Inches(0)
    abstract.paragraph_format.line_spacing = 2
    abstract.paragraph_format.space_after = Pt(0)

    reference = styles.add_style("APA Reference", 1)
    reference.base_style = normal
    set_style_font(reference)
    reference.paragraph_format.left_indent = Inches(0.5)
    reference.paragraph_format.first_line_indent = Inches(-0.5)
    reference.paragraph_format.line_spacing = 2
    reference.paragraph_format.space_after = Pt(0)

    block_quote = styles.add_style("APA Block Quote", 1)
    block_quote.base_style = normal
    set_style_font(block_quote)
    block_quote.paragraph_format.left_indent = Inches(0.5)
    block_quote.paragraph_format.first_line_indent = Inches(0)
    block_quote.paragraph_format.line_spacing = 2
    block_quote.paragraph_format.space_after = Pt(0)

    table_note = styles.add_style("APA Table or Figure Note", 1)
    table_note.base_style = normal
    set_style_font(table_note)
    table_note.paragraph_format.first_line_indent = Inches(0)
    table_note.paragraph_format.line_spacing = 2
    table_note.paragraph_format.space_after = Pt(0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.base_style = normal
        set_style_font(style, bold=True, italic=(style_name == "Heading 3"))
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Inches(0)

    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def configure_header(document):
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 2

    run = paragraph.add_run(RUNNING_HEAD)
    set_font(run)
    paragraph.add_run("\t")
    add_page_number(paragraph)


def add_paragraph(document, text="", style=None, alignment=None, bold=False, italic=False):
    paragraph = document.add_paragraph(style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    if text:
        run = paragraph.add_run(text)
        set_font(run, italic=italic, bold=bold)
    return paragraph


def add_mixed_runs(paragraph, parts):
    for text, bold, italic in parts:
        run = paragraph.add_run(text)
        set_font(run, bold=bold, italic=italic)


def add_title_page(document):
    for _ in range(3):
        add_paragraph(document)

    add_paragraph(document, FULL_TITLE, style="APA Title")
    add_paragraph(document, AUTHOR_NAMES, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, AFFILIATIONS, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(4):
        add_paragraph(document)

    add_paragraph(document, "Author Note", style="APA Centered Bold")
    add_paragraph(
        document,
        "[First paragraph: ORCID IDs or affiliation changes, if applicable.]",
        style="APA Body",
    )
    add_paragraph(
        document,
        "[Second paragraph: disclosures, funding, conflicts of interest, or acknowledgments.]",
        style="APA Body",
    )
    add_paragraph(
        document,
        "[Third paragraph: correspondence address and email for the corresponding author.]",
        style="APA Body",
    )

    document.add_page_break()


def add_abstract(document):
    add_paragraph(document, "Abstract", style="APA Centered Bold")
    add_paragraph(
        document,
        "[Write one concise paragraph summarizing the research problem, method, results, and implications. "
        "Do not indent the abstract paragraph. Most APA abstracts are no more than 250 words unless the target "
        "journal or publisher specifies another limit.]",
        style="APA Abstract",
    )
    paragraph = add_paragraph(document, style="APA Abstract")
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    add_mixed_runs(
        paragraph,
        [
            ("Keywords:", False, True),
            (" [keyword], [keyword], [keyword]", False, False),
        ],
    )
    document.add_page_break()


def add_body(document):
    add_paragraph(document, FULL_TITLE, style="APA Title")
    add_paragraph(
        document,
        "[Begin the introduction here. Do not add an 'Introduction' heading in APA style; the paper title at the "
        "top of the first text page serves that function. Include a narrative citation such as Smith (2024) and "
        "a parenthetical citation such as (Garcia & Lee, 2023).]",
        style="APA Body",
    )
    add_paragraph(
        document,
        "[Continue the introduction with the research question, brief literature context, and purpose statement. "
        "Use first-line indentation for regular body paragraphs and keep the entire manuscript double-spaced.]",
        style="APA Body",
    )

    add_paragraph(document, "Method", style="Heading 1")
    add_paragraph(
        document,
        "[Describe the design, participants, materials, and procedure. Use Level 1 headings for major sections.]",
        style="APA Body",
    )
    add_paragraph(document, "Participants", style="Heading 2")
    add_paragraph(
        document,
        "[Use Level 2 headings for subsections. Report participant characteristics and sampling procedures.]",
        style="APA Body",
    )
    add_paragraph(document, "Measures", style="Heading 2")
    add_paragraph(
        document,
        "[Describe measures or instruments. Direct quotations require a page or paragraph location, for example "
        "(Smith, 2024, p. 15).]",
        style="APA Body",
    )
    add_paragraph(document, "Scoring Procedure", style="Heading 3")
    add_paragraph(
        document,
        "[Use Level 3 headings only after Level 1 and Level 2 headings are already present.]",
        style="APA Body",
    )

    add_paragraph(document, "Results", style="Heading 1")
    add_paragraph(
        document,
        "[Summarize the analysis plan and key findings. Refer to each table and figure before it appears, as in "
        "Table 1 and Figure 1.]",
        style="APA Body",
    )

    add_table_example(document)
    add_figure_placeholder(document)

    add_paragraph(document, "Discussion", style="Heading 1")
    add_paragraph(
        document,
        "[Interpret the findings, explain limitations, and describe implications. A block quotation of 40 words "
        "or more should be indented 0.5 in from the left margin and should not use quotation marks.]",
        style="APA Body",
    )
    add_paragraph(
        document,
        "[Block quotation placeholder. Use this style only for quotations of 40 words or more. Keep the block "
        "quotation double-spaced, cite the source, and place final punctuation according to APA quotation rules.]",
        style="APA Block Quote",
    )

    document.add_page_break()


def add_table_example(document):
    add_paragraph(document, "Table 1", bold=True)
    add_paragraph(document, "Descriptive Statistics for Main Study Variables", italic=True)

    table = document.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    values = [
        ("Variable", "M", "SD"),
        ("[Variable 1]", "[0.00]", "[0.00]"),
        ("[Variable 2]", "[0.00]", "[0.00]"),
        ("[Variable 3]", "[0.00]", "[0.00]"),
    ]

    for row, row_values in zip(table.rows, values):
        for cell, value in zip(row.cells, row_values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing = 2
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_font(run, bold=(row_values == values[0]))
            set_cell_border(
                cell,
                top={"val": "nil"},
                left={"val": "nil"},
                bottom={"val": "nil"},
                right={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )

    for cell in table.rows[0].cells:
        set_cell_border(
            cell,
            top={"val": "single", "sz": "8", "space": "0", "color": "000000"},
            bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"},
        )
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "space": "0", "color": "000000"})

    note = add_paragraph(document, style="APA Table or Figure Note")
    add_mixed_runs(
        note,
        [
            ("Note.", False, True),
            (" Replace this example with a clean APA table. Explain abbreviations and table-specific notes here.", False, False),
        ],
    )


def add_figure_placeholder(document):
    add_paragraph(document)
    add_paragraph(document, "Figure 1", bold=True)
    add_paragraph(document, "Conceptual Model of the Study", italic=True)
    add_paragraph(
        document,
        "[Insert figure here. Keep image text legible and refer to the figure in the body before it appears.]",
        style="APA Table or Figure Note",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    note = add_paragraph(document, style="APA Table or Figure Note")
    add_mixed_runs(
        note,
        [
            ("Note.", False, True),
            (" Replace this placeholder with the figure and any required note, source, or copyright statement.", False, False),
        ],
    )


def add_references(document):
    add_paragraph(document, "References", style="APA Centered Bold")

    entries = [
        (
            "Author, A. A., & Author, B. B. (2024). Title of article in sentence case. "
            "Title of Periodical, 12(3), 45-67. https://doi.org/xxxxx"
        ),
        "Author, A. A. (2024). Title of book in sentence case. Publisher. https://doi.org/xxxxx",
        (
            "Group Author. (2024, Month Day). Title of webpage in sentence case. "
            "Site Name. https://www.example.com/page"
        ),
    ]

    for entry in entries:
        add_paragraph(document, entry, style="APA Reference")


def build_template():
    document = Document()
    configure_document(document)
    configure_header(document)
    add_title_page(document)
    add_abstract(document)
    add_body(document)
    add_references(document)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_template()
    print(f"Created {path}")
